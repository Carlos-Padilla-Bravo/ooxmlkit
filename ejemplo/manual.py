"""Ejemplo trabajado: el manual de identidad de Ignacia Fuentes, en Word.

    python ejemplo/manual.py

Deja ejemplo/manual.docx con el indice vacio: el indice y el total de paginas
son campos que solo Word resuelve al abrir el archivo. Para cerrarlo, cerrar.ps1.

Ignacia es ficticia y su sistema esta publicado en el repositorio hermano
identidad-personal, que entrega el mismo manual en HTML y PDF. Este archivo es
la otra salida, la editable, y ejercita de paso toda la API: portada, indice,
secciones numeradas, cuadros, una figura, una franja de color, un anexo con
letra y un pie con la paginacion.

Las decisiones de ella viven en ejemplo/ignacia.py. Este archivo solo compone.
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt

# Anclado al archivo y no al directorio de trabajo, para que corra desde
# cualquier sitio y no solo desde la raiz del repositorio.
AQUI = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(AQUI))
sys.path.insert(0, AQUI)

import ooxmlkit  # noqa: E402
from ignacia import (COND, MONO, NOTA_HOJA, NOTA_TEXTO_TER,  # noqa: E402
                     PALETA, SANS, SERIF)

# Las dos sustituciones que convierten la libreria en el sistema de una persona.
# Van antes de construir_estilos(), que lee estos valores al definir los estilos.
ooxmlkit.usar_paleta(PALETA)
ooxmlkit.SERIF, ooxmlkit.SANS = SERIF, SANS
ooxmlkit.MONO, ooxmlkit.COND = MONO, COND

from ooxmlkit import (CAJA_CM, borde_par, campo, claro_oscuro,  # noqa: E402
                      construir_estilos, contra, cuadro, figura, hsl_txt, hx,
                      idioma, marcar_anexo, muestra_color, nivel, nom_hx,
                      numeracion_de_secciones, pagina, parrafo,
                      particion_de_palabras, portada_sin_encabezado, ratio,
                      rgb, rico, sombra_par, texto)

VERSION, FECHA = "1.0", "julio de 2026"

doc = Document()
construir_estilos(doc)
particion_de_palabras(doc)
idioma(doc, "es-CL")
id_anexo = numeracion_de_secciones(doc)

sec = doc.sections[0]
pagina(sec)
portada_sin_encabezado(sec)


# ------------------------------------------------------------ encabezado y pie
pie = sec.footer.paragraphs[0]
pie.style = doc.styles["Encabezado"]
pie.paragraph_format.tab_stops.add_tab_stop(Cm(CAJA_CM), WD_TAB_ALIGNMENT.RIGHT)
borde_par(pie, "top", "borde", sz=6, espacio=6)
texto(pie, "Ignacia Fuentes · manual de identidad", fuente=COND, tam=8.5,
      color="texto_ter")
pie.add_run("\t")
r = campo(pie, " PAGE ", "1")
r.font.name, r.font.size, r.font.bold = MONO, Pt(8.5), True
r.font.color.rgb = rgb("primario")
texto(pie, " de ", fuente=COND, tam=8.5, color="texto_ter")
r = campo(pie, " NUMPAGES ", "1")
r.font.name, r.font.size = MONO, Pt(8.5)
r.font.color.rgb = rgb("texto_ter")


# --------------------------------------------------------------------- portada
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(style="PortadaMeta")
texto(p, "MANUAL DE IDENTIDAD", fuente=COND, tam=10, color="senal", negrita=True)

p = doc.add_paragraph(style="PortadaNombre")
texto(p, "Ignacia", fuente=SANS, tam=33, color="texto", negrita=True)
p = doc.add_paragraph(style="PortadaApellido")
texto(p, "Fuentes", fuente=SANS, tam=33, color="texto_sec")

barra = doc.add_paragraph()
barra.paragraph_format.space_before = Pt(10)
barra.paragraph_format.space_after = Pt(26)
barra.paragraph_format.right_indent = Cm(CAJA_CM - 2.3)
texto(barra, " ", tam=2)
sombra_par(barra, "primario")

p = doc.add_paragraph(style="PortadaTitulo")
texto(p, "Sistema de color, tipografía y voz", fuente=SANS, tam=19,
      color="texto", negrita=True)

p = doc.add_paragraph(style="Entradilla")
p.paragraph_format.right_indent = Cm(2.5)
texto(p, "Documento normativo. Define cómo se ve y cómo suena todo lo que "
         "publico, sea una minuta para una paciente o una lámina para una "
         "charla.", fuente=SERIF, tam=13, color="texto_fuerte")

for _ in range(4):
    doc.add_paragraph()

for etiqueta, valor in (("Versión", VERSION), ("Fecha", FECHA),
                        ("Compuesto con", "ooxmlkit"),
                        ("Versión en HTML", "identidad-personal")):
    p = doc.add_paragraph(style="PortadaMeta")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(4.6))
    texto(p, etiqueta.upper(), fuente=COND, tam=8.5, color="texto_ter")
    p.add_run("\t")
    texto(p, valor, fuente=MONO, tam=9.5, color="texto")


# ---------------------------------------------------------------------- índice
# Sin el estilo Heading 1, para que el índice no se liste dentro de sí mismo.
t = doc.add_paragraph()
t.paragraph_format.page_break_before = True
t.paragraph_format.space_after = Pt(10)
texto(t, "Índice", fuente=SANS, tam=23, color="texto", negrita=True)
p = doc.add_paragraph(style="Nota")
texto(p, "Vacío hasta que Word actualice los campos: seleccionarlo y pulsar F9, "
         "o ejecutar cerrar.ps1.", fuente=COND, tam=9, color="texto_ter")
doc.add_paragraph()
campo(doc.add_paragraph(), ' TOC \\o "1-2" \\h \\z \\u ',
      "Índice: actualizar los campos para generarlo.")


# ------------------------------------------------------------------ 1. Esencia
doc.add_heading("Esencia", level=1)

parrafo(doc, "La esencia es lo que queda cuando se saca todo lo demás.",
        estilo="Entradilla")

doc.add_heading("Idea central", level=2)
p = doc.add_paragraph(style="Destacado")
p.paragraph_format.left_indent = Cm(0.8)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(12)
texto(p, "Convierto la evidencia nutricional en cambios de hábito que la gente "
         "sí sostiene, no en dietas que se abandonan.",
      fuente=SERIF, tam=15, color="primario_int")
borde_par(p, "left", "senal", sz=18, espacio=10)

parrafo(doc, "Esa frase es el centro. Si le quito todo lo demás, el rubro, el "
             "tono y los colores, sigue siendo cierta. Mi trabajo no es "
             "prescribir una dieta perfecta sobre el papel, sino acompañar un "
             "cambio que aguante el día a día real de una persona: su tiempo, "
             "su presupuesto, su cocina.")

doc.add_heading("Propósito", level=2)
parrafo(doc, "Acortar la distancia entre lo que la ciencia de la nutrición ya "
             "sabe y lo que una persona hace en su cocina cada día. Hay mucho "
             "conocimiento firme que no llega, o llega deformado por la moda. "
             "Mi trabajo es ese puente.")

parrafo(doc, "Mi carácter se define eligiendo un lado, no quedándome en el "
             "medio de todo.")

cuadro(doc, 1, "Posiciones de carácter",
       ["Más", "Que", "Se nota en"],
       [["**Cercana", "distante", "Tuteo, ejemplos de la vida diaria."],
        ["**Divulgativa", "técnica", "Explico primero, el término después."],
        ["**Cálida", "sobria", "Sin culpa ni reto; acompaño, no regaño."],
        ["**Categórica con evidencia", "categórica siempre",
         "Afirmo fuerte donde hay respaldo; con la moda, cauta."]],
       anchos=[4.2, 3.4, 7.4],
       alinear=["l", "l", "j"],
       fuente_datos=[COND, COND, SERIF],
       nota="Elegir un lado no borra el otro: en una charla para colegas subo "
            "el registro técnico, pero el eje por defecto es este. Una pieza "
            "que se corra de estas marcas está fuera de la identidad, aunque "
            "por separado se vea bien.")


# ---------------------------------------------------------- 2. Sistema de color
doc.add_heading("Sistema de color", level=1)

parrafo(doc, "El color no decora: reparte jerarquía. Cada valor tiene un rol y "
             "un solo rol, y por eso se puede sustituir un color sin rehacer el "
             "documento.", estilo="Entradilla")

doc.add_heading("Los diez roles", level=2)
parrafo(doc, NOTA_HOJA)
parrafo(doc, "La franja de abajo es una tabla de una fila sin bordes, no una "
             "imagen, así que se imprime a la resolución del papel y no a la de "
             "la pantalla.")
muestra_color(doc, ["texto", "texto_fuerte", "texto_sec", "texto_ter",
                    "primario", "primario_int", "senal", "superficie"])

filas = []
for clave, uso in (("texto", "Cuerpo y títulos"),
                   ("texto_fuerte", "Entradilla, cita, tercer nivel"),
                   ("texto_sec", "Notas, pies, encabezado"),
                   ("texto_ter", "Cotas y numeración menor"),
                   ("primario", "Color de marca"),
                   ("primario_int", "Su versión intensa, para texto"),
                   ("senal", "Marcador de cuadro y figura"),
                   ("superficie", "Bloque tramado sobre la hoja"),
                   ("borde", "Filetes finos")):
    v = ratio(clave, "fondo")
    aplica = clave not in ("superficie", "borde")
    filas.append([uso, hx(clave), contra(clave),
                  nivel(v) if aplica else "no aplica"])

cuadro(doc, 2, "Rol, valor y contraste sobre la hoja",
       ["Rol", "Hex", "Contraste", "Nivel WCAG"],
       filas, anchos=[6.0, 3.2, 2.9, 2.9],
       alinear=["l", "l", "r", "l"],
       fuente_datos=[COND, MONO, MONO, COND],
       nota="Calculado con ratio() al generar el documento, no escrito a mano: "
            "si un valor cambia, este cuadro cambia solo. Superficie y borde no "
            "llevan texto encima, así que el nivel no les aplica. " +
            NOTA_TEXTO_TER)

doc.add_heading("Proporción de uso", level=2)
figura(doc, 1, os.path.join(AQUI, "proporcion.png"),
       "Cuánto de cada cosa, en una pieza cualquiera",
       nota="Diseñada a 15,5 cm e insertada a 15,0 cm. Los cuerpos de texto "
            "dentro de la figura están calibrados para esa reducción.")

doc.add_heading("Usos prohibidos", level=2)
parrafo(doc, "Delimitar por exclusión evita la discusión caso a caso. Las cuatro "
             "reglas de abajo no admiten excepción, y la primera es la que más "
             "se rompe.")
for regla in ("La terracota no ordena, señala. No sirve para numerar secciones "
              "ni para decorar un título.",
              "El verde no se aclara ni se oscurece para ganar contraste: para "
              "eso está su versión intensa.",
              "Ningún texto va sobre la crema si no alcanza 4,5 contra ella, "
              "aunque sí lo alcance contra el blanco.",
              "Nada de degradados entre el verde y la terracota: son dos roles "
              "distintos y mezclarlos borra la diferencia."):
    parrafo(doc, regla, estilo="Vineta")


# -------------------------------------------------------------- 3. Tipografía
doc.add_heading("Tipografía", level=1)

parrafo(doc, "Tres familias y ninguna más. La jerarquía la hacen el cuerpo y el "
             "peso, no la cantidad de tipos.", estilo="Entradilla")

cuadro(doc, 3, "Las tres familias y su encargo",
       ["Familia", "Dónde va", "Por qué"],
       [["**Fraunces", "Títulos, firma y citas",
         "Una serif con carácter, que le pone cara al documento sin gritar. Es "
         "la que titula, no la que se lee de corrido."],
        ["**Nunito Sans", "Cuerpo, cuadros y apoyos",
         "Una sans humanista para leer. Aguanta el párrafo largo y también el "
         "cuerpo chico de una nota al pie."],
        ["**Space Mono", "Cifras, tokens y tablas de aporte",
         "Alinea los dígitos en columna, y una cifra que se compara tiene que "
         "poder leerse en vertical."]],
       anchos=[3.6, 4.8, 6.6],
       alinear=["l", "l", "j"],
       fuente_datos=[COND, COND, SERIF],
       nota="Las tres son de licencia OFL. Word compone con lo que esté "
            "instalado y sustituye sin avisar, así que un documento que salga "
            "con otra letra se contradice a sí mismo sin que nada haya fallado.")

parrafo(doc, "Las cifras dentro del texto van en la monoespaciada, que las "
             "alinea cuando aparecen en lista.")
parrafo(doc, "12 semanas de seguimiento, 3 controles, 1 sola cosa nueva por vez.",
        estilo="Dato")

rico(doc, [("Regla de composición. ", {"negrita": True, "color": "primario_int"}),
           "El cuerpo va justificado y con partición de palabras activa, que es "
           "la combinación que evita los ríos de espacio en una medida de quince "
           "centímetros. La partición queda suprimida en las medidas cortas, "
           "donde partiría palabras dentro de una columna de tres centímetros."])


# -------------------------------------------------------------------- Anexo A
h = doc.add_heading("Ficha técnica", level=1)
marcar_anexo(h, id_anexo)

parrafo(doc, "Todo lo que hace falta para componer una pieza, en una hoja. Ante "
             "cualquier discrepancia con el resto del manual, manda este cuadro.",
        estilo="Entradilla")

cuadro(doc, 4, "Resumen del sistema",
       ["Elemento", "Valor"],
       [["**Firma",
         "Nominal, por decisión: mi cara y mi nombre son la marca, no un dibujo."],
        ["**Primario", nom_hx("Verde", "primario")],
        ["**Señal", nom_hx("Terracota", "senal") + ", racionada"],
        ["**Texto", claro_oscuro("texto", "superficie")],
        ["**Primario en HSL", hsl_txt("primario")],
        ["**Familias", "Fraunces, Nunito Sans, Space Mono"],
        ["**Caja de texto",
         f"{CAJA_CM:.1f} cm".replace(".", ",") + ", A4 con márgenes de 3 y 2,5 cm"],
        ["**Contraste mínimo", "4,5 para texto; 3,0 para elementos no textuales"]],
       anchos=[4.2, 10.8],
       alinear=["l", "j"],
       fuente_datos=[COND, SERIF],
       nota="Los valores de color no están escritos en este cuadro: salen de la "
            "paleta al generar, con hx(), nom_hx(), hsl_txt() y claro_oscuro().")

salida = os.path.join(AQUI, "manual.docx")
doc.save(salida)
print(salida)
