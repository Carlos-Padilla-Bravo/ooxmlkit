"""Ejemplo minimo: un informe corto que ejercita toda la API de ooxmlkit.

    python ejemplo/informe.py

Deja ejemplo/informe.docx con el indice vacio: el indice y el total de paginas
son campos que solo Word resuelve. Para cerrarlo, cerrar.ps1.
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt

# Anclado al archivo y no al directorio de trabajo, para que corra desde
# cualquier sitio y no solo desde la raiz del repositorio.
AQUI = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(AQUI))
from ooxmlkit import (CAJA_CM, COND, MONO, PALETA, SANS, SERIF, borde_par,
                      campo, construir_estilos, contra, cuadro, figura, hx,
                      idioma, marcar_anexo, muestra_color, nivel,
                      numeracion_de_secciones, pagina, parrafo,
                      particion_de_palabras, portada_sin_encabezado, ratio,
                      rgb, sombra_par, texto)

# Las familias por defecto son IBM Plex, de licencia libre. Si no estan
# instaladas, Word sustituye sin avisar; para usar las del sistema:
#     import ooxmlkit; ooxmlkit.SERIF = "Georgia"; ooxmlkit.SANS = "Segoe UI"

doc = Document()
construir_estilos(doc)
particion_de_palabras(doc)
idioma(doc, "es-CL")
id_anexo = numeracion_de_secciones(doc)

sec = doc.sections[0]
pagina(sec)
portada_sin_encabezado(sec)


# ------------------------------------------------------------ encabezado y pie
pie = sec.footer.paragraphs[0]
pie.style = doc.styles["Encabezado"]
pie.paragraph_format.tab_stops.add_tab_stop(Cm(CAJA_CM), WD_TAB_ALIGNMENT.RIGHT)
borde_par(pie, "top", "borde", sz=6, espacio=6)
texto(pie, "ooxmlkit · informe de ejemplo", fuente=COND, tam=8.5, color="gris")
pie.add_run("\t")
r = campo(pie, " PAGE ", "1")
r.font.name, r.font.size, r.font.bold = MONO, Pt(8.5), True
r.font.color.rgb = rgb("azul")
texto(pie, " de ", fuente=COND, tam=8.5, color="gris")
r = campo(pie, " NUMPAGES ", "1")
r.font.name, r.font.size = MONO, Pt(8.5)
r.font.color.rgb = rgb("gris")


# --------------------------------------------------------------------- portada
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(style="PortadaMeta")
texto(p, "DOCUMENTO DE EJEMPLO", fuente=COND, tam=10, color="naranja_prof",
      negrita=True)

p = doc.add_paragraph(style="PortadaNombre")
texto(p, "ooxmlkit", fuente=SANS, tam=33, color="tinta", negrita=True)

barra = doc.add_paragraph()
barra.paragraph_format.space_before = Pt(10)
barra.paragraph_format.space_after = Pt(26)
barra.paragraph_format.right_indent = Cm(CAJA_CM - 2.3)
texto(barra, " ", tam=2)
sombra_par(barra, "naranja_prof")

p = doc.add_paragraph(style="PortadaTitulo")
texto(p, "Un .docx compuesto por código", fuente=SANS, tam=19, color="tinta",
      negrita=True)

p = doc.add_paragraph(style="Entradilla")
p.paragraph_format.right_indent = Cm(2.5)
texto(p, "Este informe no tiene contenido propio: existe para mostrar qué "
         "compone la librería y cómo se lee el resultado en papel.",
      fuente=SERIF, tam=13, color="grafito")

for _ in range(4):
    doc.add_paragraph()

for etiqueta, valor in (("Versión", "1.0"), ("Generado por", "ejemplo/informe.py"),
                        ("Dependencia", "python-docx")):
    p = doc.add_paragraph(style="PortadaMeta")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(4.6))
    texto(p, etiqueta.upper(), fuente=COND, tam=8.5, color="gris")
    p.add_run("\t")
    texto(p, valor, fuente=MONO, tam=9.5, color="tinta")


# ---------------------------------------------------------------------- índice
# Sin el estilo Heading 1, para que el índice no se liste dentro de sí mismo.
t = doc.add_paragraph()
t.paragraph_format.page_break_before = True
t.paragraph_format.space_after = Pt(10)
texto(t, "Índice", fuente=SANS, tam=23, color="tinta", negrita=True)
p = doc.add_paragraph(style="Nota")
texto(p, "Vacío hasta que Word actualice los campos: seleccionarlo y pulsar F9, "
         "o ejecutar cerrar.ps1.", fuente=COND, tam=9, color="gris")
doc.add_paragraph()
campo(doc.add_paragraph(), ' TOC \\o "1-2" \\h \\z \\u ',
      "Índice: actualizar los campos para generarlo.")


# ------------------------------------------------------------- 1. Composición
doc.add_heading("Composición", level=1)

parrafo(doc, "La numeración de esta sección no está escrita en el texto. La pone "
             "Word desde una lista multinivel ligada al estilo, así que aparece "
             "en el índice, admite referencias cruzadas y se renumera sola al "
             "insertar una sección nueva.", estilo="Entradilla")

doc.add_heading("Cuerpo de texto", level=2)

parrafo(doc, "El cuerpo va justificado y con partición de palabras activa, que es "
             "la combinación que evita los ríos de espacio en una medida de quince "
             "centímetros. La partición queda suprimida en los estilos de medida "
             "corta, las notas y las celdas, donde partiría palabras dentro de una "
             "columna de tres centímetros.")

for regla in ("Los saltos de página van en el estilo, nunca en un párrafo vacío.",
              "Los títulos se mantienen con el párrafo que sigue.",
              "Las viudas y huérfanas quedan controladas en todos los estilos."):
    parrafo(doc, regla, estilo="Vineta")

doc.add_heading("Cuadros", level=2)

cuadro(doc, 1, "Lo que resuelve la librería y lo que no",
       ["Elemento", "Quién lo compone", "Cuándo"],
       [["Estilos y numeración", "**ooxmlkit", "Al generar"],
        ["Cuadros y figuras", "**ooxmlkit", "Al generar"],
        ["Índice", "**Word", "Al abrir el archivo"],
        ["Total de páginas", "**Word", "Al abrir el archivo"],
        ["Fuentes incrustadas", "**Word", "Al guardar"]],
       anchos=[5.5, 5.0, 4.5],
       fuente_datos=[SERIF, COND, COND],
       nota="Los cuadros no llevan líneas verticales y cierran con un filete "
            "grueso. Las celdas admiten saltos de línea y texto en negrita.")


# --------------------------------------------------------- 2. Figuras y datos
doc.add_heading("Figuras y datos", level=1)

parrafo(doc, "Las figuras se insertan como PNG a 300 puntos por pulgada, con el "
             "pie numerado y su nota al pie. La librería no dibuja gráficos: "
             "recibe el archivo ya compuesto.", estilo="Entradilla")

figura(doc, 1, os.path.join(AQUI, "serie.png"),
       "Una serie cualquiera, con el último dato realzado",
       nota="Diseñada a 15,5 cm e insertada a 15,0 cm. Los cuerpos de texto "
            "dentro de la figura están calibrados para esa reducción.")

parrafo(doc, "Las cifras dentro del texto van en la familia monoespaciada, que "
             "alinea los dígitos en columna cuando aparecen en lista.")
parrafo(doc, "63 unidades en 2025, contra 38 en 2021.", estilo="Dato")


# --------------------------------------------------------------- Anexo A
h = doc.add_heading("Paleta por defecto", level=1)
marcar_anexo(h, id_anexo)

parrafo(doc, "Los valores no están escritos a mano en este cuadro: salen del "
             "diccionario de la librería, y el contraste se calcula al generar.",
       estilo="Entradilla")

parrafo(doc, "La franja de abajo es una tabla de una fila sin bordes, no una "
             "imagen, así que se imprime a la resolución del papel.")
muestra_color(doc, ["tinta", "grafito", "grafito_med", "gris", "azul",
                    "azul_prof", "naranja_prof", "nieve"])

filas = []
for k in PALETA:
    if k == "blanco":
        continue
    v = ratio(k, "blanco")
    filas.append([k, hx(k), contra(k), nivel(v)])

cuadro(doc, 2, "Colores, contraste sobre blanco y nivel WCAG 2.1",
       ["Clave", "Hex", "Contraste", "Texto normal"],
       filas, anchos=[4.0, 4.0, 3.5, 3.5],
       alinear=["l", "l", "r", "l"],
       fuente_datos=[COND, MONO, MONO, COND],
       nota="Calculado con ratio(), no estimado a ojo. Nieve y borde son fondo "
            "y filete: no llevan texto encima, así que el nivel no les aplica.")

salida = os.path.join(AQUI, "informe.docx")
doc.save(salida)
print(salida)
