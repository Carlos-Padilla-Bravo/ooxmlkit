"""Capa de bajo nivel sobre python-docx: inyecta el OOXML que la libreria no cubre.

Numeracion multinivel ligada a estilos, particion de palabras, campos de Word,
bordes y sombreado de parrafo, cuadros sin lineas verticales. Lo que hace falta
para que un .docx generado por codigo se componga como un documento editorial y
no como una salida de programa.

Los nombres de la API estan en espanol, igual que el documento para el que se
escribio. Los dos README los documentan uno por uno.
"""
import colorsys

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Paleta por defecto. Toda funcion que recibe un color acepta una clave de este
# diccionario o un hexadecimal suelto, asi que se puede ignorar por completo.
# Para cambiarla, entera o en parte, usar_paleta().
PALETA = {
    "tinta":        "#1A1D21",
    "grafito":      "#33383D",
    "grafito_med":  "#63696F",
    "gris":         "#767D84",
    "azul":         "#1F5FA8",
    "azul_prof":    "#17436F",
    "naranja_prof": "#9C4A05",
    "blanco":       "#FFFFFF",
    "nieve":        "#F4F5F7",
    "borde":        "#DCDFE3",
}

# Familias tipograficas. Se reasignan desde fuera: ooxmlkit.SERIF = "Georgia".
SANS = "IBM Plex Sans"
SERIF = "IBM Plex Serif"
MONO = "IBM Plex Mono"
COND = "IBM Plex Sans Condensed"

CAJA_CM = 15.0  # ancho util de la caja de texto


def usar_paleta(colores):
    """Agrega o sustituye colores de la paleta activa.

    Solo toca las claves que se pasan, asi que sirve para cambiar un color
    suelto o para cargar una paleta propia entera.
    """
    PALETA.update(colores)
    return PALETA


def hexa(k):
    """Hexadecimal sin almohadilla de una clave de la paleta, o de un hex suelto.

    Valida el resultado. Sin esta comprobacion, una clave mal escrita entra al
    XML tal cual y Word la descarta en silencio: el documento sale mal sin que
    nada haya fallado. Se admite "auto", que OOXML acepta como color.
    """
    v = PALETA.get(k, k).lstrip("#")
    if v != "auto" and (len(v) != 6 or any(c not in "0123456789abcdefABCDEF" for c in v)):
        raise ValueError(f"color no valido: {k!r}. Debe ser una clave de PALETA "
                         f"o un hexadecimal de seis digitos.")
    return v


def rgb(k):
    h = hexa(k)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def rgb_txt(k):
    h = hexa(k)
    return ", ".join(str(int(h[i:i + 2], 16)) for i in (0, 2, 4))


def hsl_txt(k):
    h = hexa(k)
    r, g, b = (int(h[i:i + 2], 16) / 255 for i in (0, 2, 4))
    hh, ll, ss = colorsys.rgb_to_hls(r, g, b)
    return f"{round(hh * 360)}°, {round(ss * 100)} %, {round(ll * 100)} %"


def _lin(c):
    c = c / 255
    return c / 12.92 if c <= 0.04045 else ((c + 0.055) / 1.055) ** 2.4


def lum(k):
    """Luminancia relativa segun WCAG 2.1."""
    h = hexa(k)
    r, g, b = (int(h[i:i + 2], 16) for i in (0, 2, 4))
    return 0.2126 * _lin(r) + 0.7152 * _lin(g) + 0.0722 * _lin(b)


def ratio(a, b):
    """Razon de contraste WCAG 2.1 entre dos colores, de 1 a 21."""
    la, lb = lum(a), lum(b)
    hi, lo = max(la, lb), min(la, lb)
    return (hi + 0.05) / (lo + 0.05)


def nivel(v, grande=False):
    """Nivel WCAG que alcanza una razon de contraste."""
    if grande:
        return "AAA" if v >= 4.5 else "AA" if v >= 3 else "no cumple"
    return ("AAA" if v >= 7 else "AA" if v >= 4.5
            else "AA grande" if v >= 3 else "no cumple")


def contra(a, b="blanco"):
    """Razon de contraste como texto, con coma decimal."""
    return f"{ratio(a, b):.2f}".replace(".", ",")


def hx(k):
    """Hexadecimal en mayusculas, con almohadilla, de un color de la paleta."""
    return "#" + hexa(k).upper()


def claro_oscuro(k_claro, k_oscuro):
    """Declara un color en su version clara y en su version oscura.

    Ejemplo trabajado: arma una fila de la ficha tecnica de un documento
    concreto. Para otro documento se copia y se edita.
    """
    return f"{hx(k_claro)} sobre claro; {hx(k_oscuro)} sobre oscuro"


def nom_hx(nombre, k):
    """Nombre legible seguido del hexadecimal leido de la paleta."""
    return f"{nombre} {hx(k)}"


def el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


# --------------------------------------------------------------- documento
def campo(par, instruccion, texto_reserva=""):
    """Inserta un campo de Word (PAGE, TOC, NUMPAGES)."""
    r1 = par.add_run()._r
    r1.append(el("w:fldChar", **{"w:fldCharType": "begin"}))
    r2 = par.add_run()._r
    t = el("w:instrText", **{"xml:space": "preserve"})
    t.text = instruccion
    r2.append(t)
    r3 = par.add_run()._r
    r3.append(el("w:fldChar", **{"w:fldCharType": "separate"}))
    reserva = par.add_run(texto_reserva)
    r5 = par.add_run()._r
    r5.append(el("w:fldChar", **{"w:fldCharType": "end"}))
    return reserva


def actualizar_campos(doc):
    s = doc.settings.element
    s.append(el("w:updateFields", **{"w:val": "true"}))


def particion_de_palabras(doc, zona_cm=0.55):
    """Activa la particion automatica, que el texto justificado exige."""
    s = doc.settings.element
    s.append(el("w:autoHyphenation", **{"w:val": "true"}))
    s.append(el("w:doNotHyphenateCaps", **{"w:val": "true"}))
    s.append(el("w:consecutiveHyphenLimit", **{"w:val": "2"}))
    s.append(el("w:hyphenationZone", **{"w:w": int(zona_cm * 567)}))


def idioma(doc, cod="es-CL"):
    rPr = doc.styles["Normal"].element.get_or_add_rPr()
    rPr.append(el("w:lang", **{"w:val": cod, "w:eastAsia": cod, "w:bidi": "ar-SA"}))


# Sucesores de w:numPr dentro de w:pPr, segun el esquema.
_TRAS_NUMPR = (
    "w:suppressLineNumbers", "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens",
    "w:kinsoku", "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
    "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing",
    "w:ind", "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange")


def _lvl(ilvl, fmt, texto_lvl, sangria, colgado, fuente, tam, color,
         negrita=True):
    """Un nivel de lista, con el numero compuesto como elemento propio."""
    lvl = el("w:lvl", **{"w:ilvl": ilvl})
    lvl.append(el("w:start", **{"w:val": "1"}))
    lvl.append(el("w:numFmt", **{"w:val": fmt}))
    lvl.append(el("w:suff", **{"w:val": "tab"}))
    lvl.append(el("w:lvlText", **{"w:val": texto_lvl}))
    lvl.append(el("w:lvlJc", **{"w:val": "left"}))
    pPr = el("w:pPr")
    pPr.append(el("w:ind", **{"w:left": sangria, "w:hanging": colgado}))
    lvl.append(pPr)
    rPr = el("w:rPr")
    rPr.append(el("w:rFonts", **{"w:ascii": fuente, "w:hAnsi": fuente,
                                 "w:cs": fuente, "w:eastAsia": fuente}))
    if negrita:
        rPr.append(el("w:b", **{"w:val": "1"}))
    rPr.append(el("w:color", **{"w:val": hexa(color)}))
    rPr.append(el("w:sz", **{"w:val": int(tam * 2)}))
    lvl.append(rPr)
    return lvl


def numeracion_de_secciones(doc, id_sec=90, id_anexo=91):
    """Numera los titulos de nivel 1 y 2 con una lista multinivel real.

    Al ir ligada a los estilos, la numeracion aparece en el indice, admite
    referencias cruzadas y se renumera sola si se agrega una seccion.
    Devuelve el identificador de la secuencia de anexos, para marcar_anexo().
    """
    num = doc.part.numbering_part.element

    # Secciones: 1, y 1.1 para las subsecciones.
    # Sin sangria colgante: el numero arranca en el margen, igual que el
    # cuerpo de texto, y nada del documento invade los margenes.
    abst = el("w:abstractNum", **{"w:abstractNumId": id_sec})
    abst.append(el("w:multiLevelType", **{"w:val": "multilevel"}))
    abst.append(_lvl(0, "decimal", "%1", 0, 0, MONO, 20, "azul"))
    abst.append(_lvl(1, "decimal", "%1.%2", 0, 0, MONO, 14, "azul_prof"))
    for i in range(2, 9):
        abst.append(_lvl(i, "none", "", 0, 0, MONO, 11, "gris", negrita=False))

    # Anexos: letra en vez de cifra, fuera de la secuencia numerica.
    abst_a = el("w:abstractNum", **{"w:abstractNumId": id_anexo})
    abst_a.append(el("w:multiLevelType", **{"w:val": "singleLevel"}))
    abst_a.append(_lvl(0, "upperLetter", "%1", 0, 0, MONO, 20, "grafito"))

    primero = num.find(qn("w:num"))
    for a in (abst, abst_a):
        if primero is not None:
            primero.addprevious(a)
        else:
            num.append(a)
    for i in (id_sec, id_anexo):
        n = el("w:num", **{"w:numId": i})
        n.append(el("w:abstractNumId", **{"w:val": i}))
        num.append(n)

    for estilo, ilvl in (("Heading 1", 0), ("Heading 2", 1)):
        pPr = doc.styles[estilo].element.get_or_add_pPr()
        npr = el("w:numPr")
        npr.append(el("w:ilvl", **{"w:val": ilvl}))
        npr.append(el("w:numId", **{"w:val": id_sec}))
        pPr.insert_element_before(npr, *_TRAS_NUMPR)
    return id_anexo


def marcar_anexo(par, id_anexo=91):
    """Saca un titulo de la secuencia numerica y lo pasa a la de anexos."""
    pPr = par._p.get_or_add_pPr()
    viejo = pPr.find(qn("w:numPr"))
    if viejo is not None:
        pPr.remove(viejo)
    npr = el("w:numPr")
    npr.append(el("w:ilvl", **{"w:val": "0"}))
    npr.append(el("w:numId", **{"w:val": id_anexo}))
    pPr.insert_element_before(npr, *_TRAS_NUMPR)
    return par


def sin_numerar(par):
    """Quita la numeracion de un titulo suelto."""
    pPr = par._p.get_or_add_pPr()
    viejo = pPr.find(qn("w:numPr"))
    if viejo is not None:
        pPr.remove(viejo)
    npr = el("w:numPr")
    npr.append(el("w:ilvl", **{"w:val": "0"}))
    npr.append(el("w:numId", **{"w:val": "0"}))
    pPr.insert_element_before(npr, *_TRAS_NUMPR)
    return par


def pagina(seccion, izq=3.0, der=3.0, sup=2.5, inf=2.5):
    seccion.page_width = Cm(21)
    seccion.page_height = Cm(29.7)
    seccion.left_margin, seccion.right_margin = Cm(izq), Cm(der)
    seccion.top_margin, seccion.bottom_margin = Cm(sup), Cm(inf)
    seccion.header_distance, seccion.footer_distance = Cm(1.4), Cm(1.4)


def portada_sin_encabezado(seccion):
    seccion._sectPr.append(el("w:titlePg", **{"w:val": "1"}))


def borde_par(par, lado="bottom", color="borde", sz=6, espacio=4):
    pPr = par._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = el("w:pBdr")
        pPr.append(pbdr)
    pbdr.append(el(f"w:{lado}", **{"w:val": "single", "w:sz": sz,
                                   "w:space": espacio, "w:color": hexa(color)}))


def sombra_par(par, color):
    pPr = par._p.get_or_add_pPr()
    pPr.append(el("w:shd", **{"w:val": "clear", "w:color": "auto",
                              "w:fill": hexa(color)}))


# Sucesores de w:suppressAutoHyphens dentro de w:pPr, segun el esquema.
_TRAS_PARTICION = (
    "w:kinsoku", "w:wordWrap", "w:overflowPunct", "w:topLinePunct",
    "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing",
    "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection",
    "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl", "w:divId",
    "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange")


def sin_particion(pPr):
    """Desactiva la particion de palabras en un pPr ya existente."""
    if pPr.find(qn("w:suppressAutoHyphens")) is not None:
        return
    e = el("w:suppressAutoHyphens", **{"w:val": "1"})
    pPr.insert_element_before(e, *_TRAS_PARTICION)


def sin_separar(par):
    """Mantiene el parrafo con el siguiente y sin viudas ni huerfanas."""
    pf = par.paragraph_format
    pf.keep_with_next = True
    pf.widow_control = True


def salto(doc):
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------- estilos
def _estilo(doc, nombre, fuente, tam, color, negrita=False, cursiva=False,
            antes=0, despues=0, interlin=1.2, alinear=WD_ALIGN_PARAGRAPH.LEFT,
            base="Normal", mayus=False, esp_letra=None, sangria=0,
            particion=True):
    from docx.enum.style import WD_STYLE_TYPE
    try:
        st = doc.styles[nombre]
    except KeyError:
        st = doc.styles.add_style(nombre, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles[base]
    f = st.font
    f.name = fuente
    f.size = Pt(tam)
    f.bold = negrita
    f.italic = cursiva
    f.color.rgb = rgb(color)
    rPr = st.element.get_or_add_rPr()
    for attr in ("w:eastAsia", "w:cs", "w:hAnsi", "w:ascii"):
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = el("w:rFonts")
            rPr.insert(0, rf)
        rf.set(qn(attr), fuente)
    if mayus:
        rPr.append(el("w:caps", **{"w:val": "1"}))
    if esp_letra:
        rPr.append(el("w:spacing", **{"w:val": int(esp_letra * 20)}))
    pf = st.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(despues)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = interlin
    pf.alignment = alinear
    pf.left_indent = Cm(sangria)
    pf.widow_control = True
    if not particion:
        sin_particion(st.element.get_or_add_pPr())
    return st


def construir_estilos(doc):
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("tinta")
    rPr = normal.element.get_or_add_rPr()
    rf = el("w:rFonts", **{"w:ascii": SERIF, "w:hAnsi": SERIF,
                           "w:cs": SERIF, "w:eastAsia": SERIF})
    rPr.insert(0, rf)

    for h, (tam, col, antes, desp, fam) in {
            "Heading 1": (23, "tinta", 30, 10, SANS),
            "Heading 2": (16, "azul_prof", 22, 7, SANS),
            "Heading 3": (13, "grafito", 16, 5, SANS),
            "Heading 4": (11, "grafito_med", 13, 4, COND)}.items():
        st = doc.styles[h]
        st.font.name = fam
        st.font.size = Pt(tam)
        st.font.bold = True
        st.font.color.rgb = rgb(col)
        r = st.element.get_or_add_rPr()
        r.insert(0, el("w:rFonts", **{"w:ascii": fam, "w:hAnsi": fam,
                                      "w:cs": fam, "w:eastAsia": fam}))
        if h == "Heading 4":
            r.append(el("w:caps", **{"w:val": "1"}))
            r.append(el("w:spacing", **{"w:val": 12}))
        pf = st.paragraph_format
        pf.space_before, pf.space_after = Pt(antes), Pt(desp)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        pf.keep_with_next = True
        pf.widow_control = True
        if h == "Heading 1":
            # El salto va en el estilo, no en un parrafo suelto: un parrafo
            # vacio con salto puede caer al pie de la pagina y dejar una en
            # blanco antes de la seccion.
            pf.page_break_before = True

    _estilo(doc, "Cuerpo", SERIF, 11, "tinta", despues=8, interlin=1.45, alinear=J)
    _estilo(doc, "Entradilla", SERIF, 13, "grafito", despues=12, interlin=1.35,
            alinear=J)
    _estilo(doc, "Destacado", SERIF, 11, "azul_prof", despues=8, interlin=1.35,
            particion=False)
    _estilo(doc, "Nota", COND, 9, "grafito_med", despues=6, interlin=1.3,
            alinear=J, particion=False)
    _estilo(doc, "Dato", MONO, 10, "tinta", despues=6, interlin=1.35,
            particion=False)
    _estilo(doc, "TituloCuadro", COND, 9.5, "tinta", antes=12, despues=4,
            interlin=1.25, particion=False)
    _estilo(doc, "PieFigura", COND, 9.5, "tinta", antes=5, despues=4,
            interlin=1.25, particion=False)
    _estilo(doc, "NotaTabla", COND, 8.5, "grafito_med", antes=2, despues=12,
            interlin=1.25, particion=False)
    _estilo(doc, "Vineta", SERIF, 11, "tinta", despues=5, interlin=1.35,
            alinear=J, base="List Bullet", sangria=0.6)
    _estilo(doc, "Enum", SERIF, 11, "tinta", despues=5, interlin=1.35,
            alinear=J, base="List Number", sangria=0.6)
    # Las listas de regla van en bandera: no son cuerpo de texto y la
    # particion las afea en una medida corta.
    _estilo(doc, "Regla", COND, 10, "tinta", despues=5, interlin=1.3,
            sangria=0.6, particion=False)
    _estilo(doc, "Imagen", SERIF, 11, "tinta", antes=10, despues=0,
            particion=False)
    _estilo(doc, "PortadaNombre", SANS, 33, "tinta", despues=0, interlin=1.0, particion=False)
    _estilo(doc, "PortadaApellido", SANS, 33, "grafito_med", despues=0, interlin=1.0, particion=False)
    _estilo(doc, "PortadaTitulo", SANS, 19, "tinta", antes=0, despues=6, interlin=1.15, particion=False)
    _estilo(doc, "PortadaMeta", MONO, 9.5, "grafito_med", despues=3, interlin=1.4, particion=False)
    _estilo(doc, "Encabezado", COND, 8.5, "grafito_med", despues=0,
            interlin=1.0, particion=False)
    _estilo(doc, "Cita", COND, 10, "grafito", despues=8, interlin=1.35,
            alinear=J, sangria=0.8)


def texto(par, s, fuente=None, tam=None, color=None, negrita=None, cursiva=None):
    r = par.add_run(s)
    if fuente:
        r.font.name = fuente
        rPr = r._r.get_or_add_rPr()
        rPr.insert(0, el("w:rFonts", **{"w:ascii": fuente, "w:hAnsi": fuente,
                                        "w:cs": fuente, "w:eastAsia": fuente}))
    if tam:
        r.font.size = Pt(tam)
    if color:
        r.font.color.rgb = rgb(color)
    if negrita is not None:
        r.font.bold = negrita
    if cursiva is not None:
        r.font.italic = cursiva
    return r


def parrafo(doc, s="", estilo="Cuerpo"):
    p = doc.add_paragraph(style=estilo)
    if s:
        p.add_run(s)
    return p


def rico(doc, partes, estilo="Cuerpo"):
    """partes: lista de (texto, dict de formato) o str."""
    p = doc.add_paragraph(style=estilo)
    for parte in partes:
        if isinstance(parte, str):
            p.add_run(parte)
        else:
            texto(p, parte[0], **parte[1])
    return p


# --------------------------------------------------------------- tablas
def _celda_borde(celda, lado, sz, color, val="single"):
    tcPr = celda._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = el("w:tcBorders")
        tcPr.append(tb)
    viejo = tb.find(qn(f"w:{lado}"))
    if viejo is not None:
        tb.remove(viejo)
    tb.append(el(f"w:{lado}", **{"w:val": val, "w:sz": sz, "w:space": 0,
                                 "w:color": hexa(color)}))


def _celda_fondo(celda, color):
    celda._tc.get_or_add_tcPr().append(
        el("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hexa(color)}))


def _celda_margen(celda, sup=0.10, inf=0.10, izq=0.15, der=0.15):
    tcPr = celda._tc.get_or_add_tcPr()
    m = el("w:tcMar")
    for lado, v in (("top", sup), ("bottom", inf), ("start", izq), ("end", der)):
        m.append(el(f"w:{lado}", **{"w:w": int(v * 567), "w:type": "dxa"}))
    tcPr.append(m)


def cuadro(doc, n, titulo, encabezados, filas, anchos, nota=None,
           alinear=None, fuente_datos=None):
    """Cuadro numerado: titulo arriba, nota al pie. Sin lineas verticales."""
    p = doc.add_paragraph(style="TituloCuadro")
    texto(p, f"Cuadro {n}. ", negrita=True, color="naranja_prof")
    texto(p, titulo)
    sin_separar(p)

    t = doc.add_table(rows=1, cols=len(encabezados))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    tblPr = t._tbl.tblPr
    tblPr.append(el("w:tblLayout", **{"w:type": "fixed"}))
    alinear = alinear or ["l"] * len(encabezados)
    fuente_datos = fuente_datos or [SERIF] * len(encabezados)
    mapa = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT, "j": WD_ALIGN_PARAGRAPH.JUSTIFY}

    enc = t.rows[0]
    trPr = enc._tr.get_or_add_trPr()
    trPr.append(el("w:tblHeader"))
    trPr.append(el("w:cantSplit"))
    for i, (h, a) in enumerate(zip(encabezados, alinear)):
        c = enc.cells[i]
        c.width = Cm(anchos[i])
        c.text = ""
        par = c.paragraphs[0]
        par.alignment = mapa[a]
        par.paragraph_format.space_before = Pt(3)
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.line_spacing = 1.15
        texto(par, h, fuente=COND, tam=9, color="tinta", negrita=True)
        _celda_fondo(c, "nieve")
        _celda_margen(c)
        _celda_borde(c, "top", 16, "tinta")
        _celda_borde(c, "bottom", 8, "tinta")
        _celda_borde(c, "left", 0, "borde", "none")
        _celda_borde(c, "right", 0, "borde", "none")

    for j, fila in enumerate(filas):
        r = t.add_row()
        r._tr.get_or_add_trPr().append(el("w:cantSplit"))
        ultima = j == len(filas) - 1
        for i, val in enumerate(fila):
            c = r.cells[i]
            c.width = Cm(anchos[i])
            c.text = ""
            par = c.paragraphs[0]
            par.alignment = mapa[alinear[i]]
            par.paragraph_format.space_before = Pt(3)
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.line_spacing = 1.2
            fam = fuente_datos[i]
            for k, linea in enumerate(str(val).split("\n")):
                if k:
                    par = c.add_paragraph()
                    par.alignment = mapa[alinear[i]]
                    par.paragraph_format.space_before = Pt(0)
                    par.paragraph_format.space_after = Pt(3)
                    par.paragraph_format.line_spacing = 1.2
                sin_particion(par._p.get_or_add_pPr())
                neg = linea.startswith("**")
                linea = linea.replace("**", "")
                texto(par, linea, fuente=fam, tam=9 if fam != SERIF else 9.5,
                      color="tinta" if neg else "grafito", negrita=neg)
            _celda_margen(c)
            _celda_borde(c, "bottom", 16 if ultima else 3,
                         "tinta" if ultima else "borde")
            _celda_borde(c, "left", 0, "borde", "none")
            _celda_borde(c, "right", 0, "borde", "none")
            _celda_borde(c, "top", 0, "borde", "none")

    if nota:
        p = doc.add_paragraph(style="NotaTabla")
        texto(p, "Nota. ", negrita=True)
        texto(p, nota)
    else:
        doc.add_paragraph(style="NotaTabla")
    return t


def figura(doc, n, archivo, titulo, nota=None, ancho=CAJA_CM):
    """Figura numerada, con pie y nota, insertada al ancho dado en centimetros."""
    p = doc.add_paragraph(style="Imagen")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(archivo, width=Cm(ancho))
    cap = doc.add_paragraph(style="PieFigura")
    cap.paragraph_format.keep_with_next = bool(nota)
    texto(cap, f"Figura {n}. ", negrita=True, color="naranja_prof")
    texto(cap, titulo)
    if nota:
        q = doc.add_paragraph(style="NotaTabla")
        texto(q, "Nota. ", negrita=True)
        texto(q, nota)
    else:
        doc.add_paragraph(style="NotaTabla")


def muestra_color(doc, claves, alto_cm=0.55):
    """Franja de color solida dentro del documento, usando una tabla de 1 fila."""
    t = doc.add_table(rows=1, cols=len(claves))
    t.autofit = False
    ancho = CAJA_CM / len(claves)
    for i, k in enumerate(claves):
        c = t.rows[0].cells[i]
        c.width = Cm(ancho)
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_before = Pt(alto_cm * 28.35 / 2)
        par.paragraph_format.space_after = Pt(alto_cm * 28.35 / 2)
        _celda_fondo(c, k)
        for lado in ("top", "bottom", "left", "right"):
            _celda_borde(c, lado, 0, "borde", "none")
    return t
